import docx
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from ollama import Client

class TestOllama:

    def __init__(self):
        self.input_file_path = None
        self.output_file_path = None
        self.target_language = None
        self.ollama_client = Client(host="http://127.0.0.1:11434")
        self.chunk_size = 2000  # Adjust chunk size based on LLM's token limit

    def translate(self, text_to_translate: str):
        ollama_response = self.ollama_client.generate(
            model="llama3",
            prompt=f"Translate the following text into {self.target_language}:\n\n{text_to_translate}"
        )
        return ollama_response['response'].lstrip()

    def read_docx(self, file_path):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return full_text

    def write_docx(self, paragraphs, file_path):
        doc = docx.Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(file_path)

    def chunk_text(self, paragraphs, chunk_size):
        chunks = []
        current_chunk = ""
        current_size = 0

        for para in paragraphs:
            if current_size + len(para.split()) > chunk_size:
                chunks.append(current_chunk)
                current_chunk = para
                current_size = len(para.split())
            else:
                if current_chunk:
                    current_chunk += "\n" + para
                else:
                    current_chunk = para
                current_size += len(para.split())

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def translate_document(self, paragraphs):
        chunks = self.chunk_text(paragraphs, self.chunk_size)
        translated_chunks = []
        for chunk in chunks:
            translated_chunk = self.translate(chunk)
            translated_chunks.append(translated_chunk)
        translated_paragraphs = "\n".join(translated_chunks).split("\n")
        return translated_paragraphs

    def select_input_file(self):
        root = tk.Tk()
        root.withdraw()  # Hide the main window

        # Allow user to select input file
        self.input_file_path = filedialog.askopenfilename(title="Select Input File", filetypes=[("Word files", "*.docx")])

    def select_output_file(self):
        root = tk.Tk()
        root.withdraw()  # Hide the main window

        # Allow user to select output directory
        output_dir = filedialog.askdirectory(title="Select Output Directory")
        if output_dir:
            # Allow user to specify output filename
            output_filename = filedialog.asksaveasfilename(title="Save As", initialdir=output_dir, defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if output_filename:
                self.output_file_path = output_filename

    def run_gui(self):
        root = tk.Tk()
        root.title("Local Llama3 Document Translator")

        # Set window size
        root.geometry("500x200")

        # Target language dropdown
        tk.Label(root, text="Select Target Language").grid(row=0, column=0, padx=10, pady=10)
        target_language_var = tk.StringVar(root)
        target_language_options = ["Traditional Chinese", "Simplified Chinese", "English", "Japanese", "Spanish", "French", "German", "Italian", "Portuguese", "Dutch", "Russian", "Korean"]
        target_language_menu = ttk.Combobox(root, textvariable=target_language_var, values=target_language_options)
        target_language_menu.grid(row=0, column=1, padx=10, pady=10)

        # Select input file button
        select_input_button = tk.Button(root, text="Select Input File", command=self.select_input_file)
        select_input_button.grid(row=1, column=0, columnspan=2, pady=10)

        # Select output file button
        select_output_button = tk.Button(root, text="Select Output File", command=self.select_output_file)
        select_output_button.grid(row=2, column=0, columnspan=2, pady=10)

        # Start translation button
        def start_translation():
            self.target_language = target_language_var.get()
            if self.input_file_path and self.output_file_path and self.target_language:
                # Read the text from the input file
                original_paragraphs = self.read_docx(self.input_file_path)

                # Translate the text
                translated_paragraphs = self.translate_document(original_paragraphs)

                # Write the translated text to the output file
                self.write_docx(translated_paragraphs, self.output_file_path)

                messagebox.showinfo("Success", f"Translation completed. Translated document saved to {self.output_file_path}")
            else:
                messagebox.showerror("Error", "Please select all fields")

        start_translation_button = tk.Button(root, text="Start Translation", command=start_translation)
        start_translation_button.grid(row=3, column=0, columnspan=2, pady=10)

        root.mainloop()

if __name__ == '__main__':
    job = TestOllama()
    job.run_gui()
